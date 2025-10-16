# data_generation.py
# ---------------------------------------
# This script creates 7000 random Word documents (.docx)
# Each file contains random text, headings, and short paragraphs
# to simulate real-world academic or technical reports.

import os
import random
from docx import Document
from tqdm import tqdm

def generate_docx_files(n=7000, output_dir="data/source_documents"):
    """Generate n random Word documents with different content."""
    os.makedirs(output_dir, exist_ok=True)
    topics = ["Artificial Intelligence", "Biology", "History", "Mathematics", "Economics", "Physics"]

    for i in tqdm(range(n), desc="Generating DOCX files"):
        doc = Document()
        doc.add_heading(f"Report {i + 1}", level=1)

        # Add random number of sections
        for _ in range(random.randint(3, 7)):
            topic = random.choice(topics)
            doc.add_heading(topic, level=2)
            doc.add_paragraph(
                f"This section discusses the topic of {topic.lower()} in detail. "
                f"Further analysis explores the theoretical and practical aspects of {topic.lower()}."
            )

        # Occasionally add a bullet list
        if random.random() > 0.5:
            doc.add_heading("Key Points", level=2)
            for j in range(random.randint(3, 5)):
                doc.add_paragraph(f"Point {j + 1}: Insight about {random.choice(topics)}.", style="List Bullet")

        doc.save(os.path.join(output_dir, f"doc_{i + 1}.docx"))

if __name__ == "__main__":
    generate_docx_files()
