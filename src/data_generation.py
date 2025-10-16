import os
import random
from docx import Document
from tqdm import tqdm


def generate_docx_files(n=7000, output_dir="data/source_documents"):
os.makedirs(output_dir, exist_ok=True)
topics = ["AI", "Biology", "History", "Math", "Economics"]


for i in tqdm(range(n), desc="Generating DOCX files"):
doc = Document()
doc.add_heading(f"Report {i+1}", level=1)
for _ in range(random.randint(3, 7)):
topic = random.choice(topics)
doc.add_paragraph(f"This section discusses {topic} in detail.")
doc.save(os.path.join(output_dir, f"doc_{i+1}.docx"))


if __name__ == "__main__":
generate_docx_files()
